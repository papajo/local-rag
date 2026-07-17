"""Document ingestion pipeline — walk, extract, chunk, embed, store."""

from __future__ import annotations

import hashlib
import re
import shutil
import sqlite3
import subprocess
import tempfile
from pathlib import Path

import tiktoken
from rich.console import Console
from rich.progress import Progress

from local_rag.config import (
    INGEST_STATE_DB,
    MAX_FILE_SIZE_BYTES,
    MAX_SCAN_DEPTH,
    SUPPORTED_EXTENSIONS,
    get_settings,
)
from local_rag.metadata import extract_file_metadata
from local_rag.models import DocumentChunk
from local_rag.store import chunk_id, delete_file, find_similar_chunk, store_chunks, update_chunk_source_refs

# ── Tokenizer ────────────────────────────────────────────────────────────────
_tokenizer = tiktoken.get_encoding("cl100k_base")


def count_tokens(text: str) -> int:
    return len(_tokenizer.encode(text))


def _supported(path: Path) -> bool:
    return path.suffix.lower() in SUPPORTED_EXTENSIONS


# ── File hash (idempotency) ──────────────────────────────────────────────────


def file_hash(path: Path) -> str:
    """SHA-256 of file contents."""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        while True:
            chunk = f.read(65536)
            if not chunk:
                break
            h.update(chunk)
    return h.hexdigest()


# ── Ingest state DB ─────────────────────────────────────────────────────────


def _get_state_db() -> sqlite3.Connection:
    db_path = INGEST_STATE_DB
    db_path.parent.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(str(db_path))
    conn.execute(
        "CREATE TABLE IF NOT EXISTS ingest_state ("
        "  file_path TEXT PRIMARY KEY,"
        "  file_hash TEXT NOT NULL,"
        "  last_ingested TEXT DEFAULT (datetime('now'))"
        ")"
    )
    conn.commit()
    return conn


def _is_unchanged(file_path: str, fhash: str) -> bool:
    conn = _get_state_db()
    row = conn.execute(
        "SELECT file_hash FROM ingest_state WHERE file_path = ?", (file_path,)
    ).fetchone()
    conn.close()
    return row is not None and row[0] == fhash


def _mark_ingested(file_path: str, fhash: str) -> None:
    conn = _get_state_db()
    conn.execute(
        "INSERT OR REPLACE INTO ingest_state (file_path, file_hash) VALUES (?, ?)",
        (file_path, fhash),
    )
    conn.commit()
    conn.close()


# ── Generation check (LanceDB / ingest_state desync detection) ────────────────


def _ensure_metadata(conn: sqlite3.Connection) -> None:
    conn.execute(
        "CREATE TABLE IF NOT EXISTS ingest_metadata "
        "(key TEXT PRIMARY KEY, value TEXT NOT NULL)"
    )
    conn.commit()


def _get_metadata_value(conn: sqlite3.Connection, key: str, default: str = "") -> str:
    row = conn.execute(
        "SELECT value FROM ingest_metadata WHERE key = ?", (key,)
    ).fetchone()
    return row[0] if row else default


def _set_metadata_value(conn: sqlite3.Connection, key: str, value: str) -> None:
    conn.execute(
        "INSERT OR REPLACE INTO ingest_metadata (key, value) VALUES (?, ?)",
        (key, value),
    )
    conn.commit()


def _check_generation(console: Console) -> None:
    from local_rag.config import DEFAULT_DATA_DIR
    from local_rag.store import write_generation

    gen_file = DEFAULT_DATA_DIR / "rag_generation"
    if gen_file.exists():
        try:
            lancedb_gen = int(gen_file.read_text().strip())
        except (ValueError, OSError):
            lancedb_gen = 0
    else:
        lancedb_gen = 0

    conn = _get_state_db()
    _ensure_metadata(conn)
    raw = _get_metadata_value(conn, "rag_generation", "0")
    try:
        state_gen = int(raw)
    except (ValueError, TypeError):
        state_gen = 0

    if lancedb_gen != state_gen:
        console.print(
            f"[yellow]Generation mismatch (LanceDB gen={lancedb_gen}, "
            f"state gen={state_gen}). Resetting ingest state — "
            f"all files will be re-indexed.[/]"
        )
        conn.execute("DELETE FROM ingest_state")

        # Re-sync: update both the file AND state metadata to the current
        # lancedb_gen so there isn't a lingering mismatch on the next run.
        if not gen_file.exists():
            write_generation(lancedb_gen)
        _set_metadata_value(conn, "rag_generation", str(lancedb_gen))

    conn.close()


def _sync_generation() -> None:
    from local_rag.store import read_generation

    gen = read_generation()
    if gen == 0:
        return
    conn = _get_state_db()
    _ensure_metadata(conn)
    _set_metadata_value(conn, "rag_generation", str(gen))
    conn.close()


# ── Text extraction ──────────────────────────────────────────────────────────


def extract_text(filepath: Path) -> str:
    """Extract text content from a supported file type."""
    suffix = filepath.suffix.lower()

    if suffix == ".txt":
        return filepath.read_text(encoding="utf-8", errors="replace")

    if suffix == ".md":
        return filepath.read_text(encoding="utf-8", errors="replace")

    # Code and config files — read as plain text
    if suffix in (
        ".py", ".pyi", ".rs", ".ts", ".tsx", ".js", ".go",
        ".java", ".yaml", ".yml", ".toml", ".json", ".sql",
    ):
        return filepath.read_text(encoding="utf-8", errors="replace")

    if suffix == ".pdf":
        return _extract_pdf(filepath)

    if suffix == ".docx":
        return _extract_docx(filepath)

    return ""


def _extract_pdf(filepath: Path) -> str:
    """Extract text from PDF using pypdf."""
    try:
        import pypdf
    except ImportError:
        print("  [yellow]pypdf not installed, falling back to pypdfium2[/]")
        return _extract_pdf_fallback(filepath)

    try:
        reader = pypdf.PdfReader(str(filepath))
        pages: list[str] = []
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                pages.append(text.strip())
        return "\n\n".join(pages)
    except Exception as e:
        print(f"  [red]pypdf error for {filepath.name}: {e}, trying fallback[/]")
        return _extract_pdf_fallback(filepath)


def _extract_pdf_fallback(filepath: Path) -> str:
    """Fallback PDF extraction using pypdfium2."""
    try:
        import pypdfium2 as pdfium

        pdf = pdfium.PdfDocument(str(filepath))
        pages: list[str] = []
        for i in range(len(pdf)):
            textpage = pdf[i].get_textpage()
            text = textpage.get_text_range()
            if text and text.strip():
                pages.append(text.strip())
        pdf.close()
        return "\n\n".join(pages)
    except ImportError:
        print(f"  [red]No PDF library available, skipping {filepath.name}[/]")
        return ""


def _extract_docx(filepath: Path) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document

        doc = Document(str(filepath))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    except Exception as e:
        print(f"  [red]docx error for {filepath.name}: {e}[/]")
        return ""


# ── Smarter code chunking ─────────────────────────────────────────────────────

# Regex patterns for detecting top-level declaration boundaries across languages.
# Each pattern matches a line that starts a new function/class/method/etc.
_CODE_BOUNDARY_PATTERNS: dict[str, list[str]] = {
    "python": [
        r"^\s*(class\s+\w+)",
        r"^\s*(async\s+def\s+\w+)",
        r"^\s*(def\s+\w+)",
        r"^\s*(@\w+)",  # decorators
    ],
    "rust": [
        r"^\s*(fn\s+\w+)",
        r"^\s*(struct\s+\w+)",
        r"^\s*(enum\s+\w+)",
        r"^\s*(trait\s+\w+)",
        r"^\s*(impl\s+)",
        r"^\s*(mod\s+\w+)",
        r"^\s*(type\s+\w+)",
        r"^\s*(const\s+\w+)",
        r"^\s*(macro_rules!\s+\w+)",
    ],
    "typescript": [
        r"^\s*(export\s+(default\s+)?(function|class|interface|type|enum|const|let|var)\s+\w+)",
        r"^\s*(function\s+\w+)",
        r"^\s*(class\s+\w+)",
        r"^\s*(interface\s+\w+)",
        r"^\s*(type\s+\w+)",
        r"^\s*(enum\s+\w+)",
        r"^\s*(const\s+\w+\s*[=:])",
    ],
    "go": [
        r"^\s*(func\s+\w+)",
        r"^\s*(type\s+\w+\s+(struct|interface))",
        r"^\s*(const\s+\()",
        r"^\s*(var\s+\()",
    ],
    "java": [
        r"^\s*(public|private|protected|static|final|abstract|synchronized|\s)*(class|interface|enum|@interface)\s+\w+",
        r"^\s*(public|private|protected|static|final|synchronized|\s)*\w+\s+\w+\s*\(",
    ],
    "sql": [
        r"^\s*(CREATE\s+(OR\s+REPLACE\s+)?(TABLE|VIEW|FUNCTION|PROCEDURE|INDEX|TRIGGER|EVENT|TYPE)\s+\w+)",
        r"^\s*(ALTER\s+(TABLE|VIEW|FUNCTION|PROCEDURE|INDEX)\s+\w+)",
        r"^\s*(DROP\s+(TABLE|VIEW|FUNCTION|PROCEDURE|INDEX)\s+\w+)",
    ],
}


def _detect_code_lang(source: str) -> str | None:
    """Detect programming language from filename extension."""
    ext = source.rsplit(".", 1)[-1].lower() if "." in source else ""
    lang_map = {
        "py": "python", "pyi": "python",
        "rs": "rust",
        "ts": "typescript", "tsx": "typescript", "js": "typescript",
        "go": "go",
        "java": "java",
        "sql": "sql",
    }
    return lang_map.get(ext)


def _chunk_code(text: str, source: str, max_tokens: int) -> list[tuple[str, int, int]]:
    """Chunk source code by function/class boundaries.

    Splits at top-level declarations (def, class, fn, func, etc.) and
    groups them into chunks that stay under max_tokens.  Falls back to
    token-based splitting for individual declarations that exceed the limit.
    """
    import re

    lang = _detect_code_lang(source)
    patterns = _CODE_BOUNDARY_PATTERNS.get(lang, [])
    if not patterns:
        return []

    combined = "|".join(f"(?:{p})" for p in patterns)
    regex = re.compile(combined, re.MULTILINE)

    boundaries: list[int] = []
    for match in regex.finditer(text):
        boundaries.append(match.start())

    if not boundaries:
        return []

    segments: list[str] = []
    for i, start_pos in enumerate(boundaries):
        end_pos = boundaries[i + 1] if i + 1 < len(boundaries) else len(text)
        segment = text[start_pos:end_pos].strip()
        if segment:
            segments.append(segment)

    if not segments:
        return []

    chunks: list[tuple[str, int, int]] = []
    current_parts: list[str] = []
    current_tokens = 0

    for seg in segments:
        seg_tokens = count_tokens(seg)

        if seg_tokens > max_tokens:
            if current_parts:
                combined_text = "\n\n".join(current_parts)
                chunks.append((combined_text, len(chunks), 0))
                current_parts = []
                current_tokens = 0
            seg_tokens_list = _tokenizer.encode(seg)
            start = 0
            while start < len(seg_tokens_list):
                end = min(start + max_tokens, len(seg_tokens_list))
                sub_text = _tokenizer.decode(seg_tokens_list[start:end]).strip()
                if sub_text:
                    chunks.append((sub_text, len(chunks), 0))
                start += max_tokens - get_settings().chunk_overlap_tokens
            continue

        if current_tokens + seg_tokens > max_tokens and current_parts:
            combined_text = "\n\n".join(current_parts)
            chunks.append((combined_text, len(chunks), 0))
            current_parts = []
            current_tokens = 0

        current_parts.append(seg)
        current_tokens += seg_tokens

    if current_parts:
        combined_text = "\n\n".join(current_parts)
        chunks.append((combined_text, len(chunks), 0))

    total = len(chunks)
    return [(t, i, total) for (t, i, _) in chunks]


# ── General chunking ─────────────────────────────────────────────────────────


def chunk_text(text: str, source: str) -> list[tuple[str, int, int]]:
    """Sliding-window chunking with sentence-boundary snapping.

    For code files (.py, .rs, .ts, etc.), dispatches to _chunk_code() which
    splits by function/class boundaries instead.

    Returns list of (text, chunk_index, total_chunks).
    """
    settings = get_settings()
    max_tokens = settings.chunk_max_tokens
    overlap = settings.chunk_overlap_tokens

    # Handle empty text
    if not text.strip():
        return []

    if _detect_code_lang(source):
        code_chunks = _chunk_code(text, source, max_tokens)
        if code_chunks:
            return code_chunks

    tokens = _tokenizer.encode(text)
    total_tokens = len(tokens)

    if total_tokens <= max_tokens:
        return [(text.strip(), 0, 1)]

    chunks: list[tuple[str, int, int]] = []
    start = 0

    while start < total_tokens:
        end = min(start + max_tokens, total_tokens)
        chunk_tokens = tokens[start:end]
        chunk_text_str = _tokenizer.decode(chunk_tokens)

        # Sentence-boundary snapping: try to end at a sentence boundary
        if end < total_tokens:
            snapped = _snap_to_sentence_end(chunk_text_str)
            if snapped != chunk_text_str:
                chunk_text_str = snapped
                snapped_tokens = _tokenizer.encode(chunk_text_str)
                end = start + len(snapped_tokens)

        chunks.append((chunk_text_str.strip(), len(chunks), 0))
        start += max_tokens - overlap

    total = len(chunks)
    result = [(t, i, total) for (t, i, _) in chunks]
    return result


def _snap_to_sentence_end(text: str) -> str:
    """Snap to last sentence-ending punctuation within the text."""
    for sep in ["\n\n", "\n", ". ", "! ", "? "]:
        idx = text.rfind(sep)
        if idx > len(text) * 0.5:
            end = idx + len(sep.rstrip())
            return text[:end]
    return text


# ── Embedding ────────────────────────────────────────────────────────────────


def embed_texts(texts: list[str], batch_size: int = 64) -> list[list[float]]:
    """Embed texts via SIE /v1/embeddings in batches to avoid server overload."""
    import httpx

    settings = get_settings()
    url = f"{settings.sie_base_url}/v1/embeddings"

    all_embeddings: list[list[float]] = []
    for i in range(0, len(texts), batch_size):
        batch = texts[i : i + batch_size]
        payload = {"model": settings.embed_model, "input": batch}
        resp = httpx.post(url, json=payload, timeout=120.0)
        resp.raise_for_status()
        data = resp.json()
        sorted_data = sorted(data["data"], key=lambda x: x["index"])
        all_embeddings.extend(item["embedding"] for item in sorted_data)

    return all_embeddings


# ── Walk documents ───────────────────────────────────────────────────────────


# Directories whose entire subtree should be pruned during walks.
# Hidden dirs (.venv, .git, etc.) are also skipped automatically.
SKIP_DIRS = frozenset({
    "venv", "node_modules", "__pycache__",
    ".svn", ".hg", ".DS_Store", ".tox", ".mypy_cache", ".ruff_cache",
    ".pytest_cache", ".eggs", "egg-info", ".git-svn",
    "site-packages", "Pods", ".next", ".cache", ".gradle",
    ".bundle", "vendor", ".terraform", ".serverless",
})


def _is_skip_dir(name: str) -> bool:
    """Return True if this directory name should be skipped."""
    if name.startswith(".") or name in SKIP_DIRS:
        return True
    if name.endswith((".egg-info", ".dist-info")):
        return True
    return False


def walk_docs(
    paths: list[Path] | None = None,
    extra_skip_dirs: set[str] | None = None,
) -> list[Path]:
    """Recursively find supported documents under given paths.

    Uses an iterative BFS with directory pruning so skipped dirs (node_modules,
    .venv, __pycache__, etc.) are never descended into.  Also enforces a max
    scan depth and file size limit to avoid tool caches and binary blobs.
    """
    settings = get_settings()
    if not paths:
        paths = settings.scan_dirs

    found: list[Path] = []
    for p in paths:
        p = p.expanduser().resolve()
        if not p.exists():
            print(f"  [yellow]Path does not exist: {p}[/]")
            continue
        if p.is_file():
            if _supported(p):
                found.append(p)
            continue

        # BFS with pruning — each entry is (path, depth)
        queue: list[tuple[Path, int]] = [(p, 0)]
        while queue:
            root, depth = queue.pop(0)

            try:
                entries = list(root.iterdir())
            except PermissionError:
                continue

            dirs: list[Path] = []
            files: list[Path] = []
            for entry in entries:
                if entry.is_dir():
                    dirs.append(entry)
                elif entry.is_file():
                    files.append(entry)

            # Prune skipped dirs and depth violations
            if depth < MAX_SCAN_DEPTH:
                for d in dirs:
                    if _is_skip_dir(d.name):
                        continue
                    if extra_skip_dirs and d.name in extra_skip_dirs:
                        continue
                    queue.append((d, depth + 1))

            for f in files:
                if not _supported(f):
                    continue
                try:
                    if f.stat().st_size > MAX_FILE_SIZE_BYTES:
                        continue
                except OSError:
                    continue
                found.append(f)

    return sorted(found)


# ── Main ingest orchestrator ────────────────────────────────────────────────


def ingest(paths: list[Path] | None = None) -> dict:
    """Run the full ingestion pipeline.

    Returns summary dict with counts.
    """
    from rich.console import Console

    console = Console()
    settings = get_settings()

    # Ensure LanceDB table exists BEFORE the generation check so that
    # read_generation() returns the real file value, not 0 (which happens
    # when chunks.lance doesn't exist yet).
    from local_rag.store import _get_db, _get_table
    _get_table(_get_db(), create=True)

    _check_generation(console)

    docs = walk_docs(paths)
    if not docs:
        console.print("[yellow]No documents found to ingest.[/]")
        return {"found": 0, "new": 0, "skipped": 0, "failed": 0, "chunks": 0}

    console.print(f"[bold]Found {len(docs)} document(s)[/]")

    new_files: list[Path] = []
    skipped = 0
    failed = 0
    total_chunks = 0

    # Check what's new
    for doc in docs:
        fhash = file_hash(doc)
        if not _is_unchanged(str(doc), fhash):
            new_files.append(doc)
        else:
            skipped += 1

    if not new_files:
        console.print("[green]All files already indexed and unchanged.[/]")
        return {
            "found": len(docs),
            "new": 0,
            "skipped": skipped,
            "failed": 0,
            "chunks": 0,
        }

    console.print(f"[bold]{len(new_files)} new/changed file(s) to process[/]")

    # Process each file
    with Progress() as progress:
        task = progress.add_task("Ingesting...", total=len(new_files))

        for doc in new_files:
            fhash = file_hash(doc)
            try:
                progress.console.print(f"  [dim]Extracting: {doc.name}[/]")
                text = extract_text(doc)

                if not text.strip():
                    progress.console.print(f"  [yellow]  No text extracted, skipping[/]")
                    failed += 1
                    progress.advance(task)
                    continue

                # Chunk
                chunks_data = chunk_text(text, doc.name)
                if not chunks_data:
                    progress.console.print(f"  [yellow]  No chunks produced, skipping[/]")
                    failed += 1
                    progress.advance(task)
                    continue

                # Embed all chunks
                texts_to_embed = [c[0] for c in chunks_data]
                progress.console.print(f"  [dim]  Embedding {len(texts_to_embed)} chunk(s)...[/]")
                embeddings = embed_texts(texts_to_embed)

                dedup_threshold = settings.dedup_threshold
                file_metadata = extract_file_metadata(str(doc), str(settings.scan_dirs[0]))
                document_chunks: list[DocumentChunk] = []
                dedup_refs: list[dict] = []

                for (text_str, ci, total), emb in zip(chunks_data, embeddings):
                    dc = DocumentChunk(
                        id=chunk_id(fhash, ci),
                        text=text_str,
                        source=doc.name,
                        file_path=str(doc),
                        chunk_index=ci,
                        total_chunks=total,
                        file_hash=fhash,
                        embedding=emb,
                        metadata=file_metadata,
                    )
                    # Check for semantic duplicates against existing chunks
                    match = find_similar_chunk(emb, dedup_threshold, exclude_hash=fhash)
                    if match:
                        new_ref = {
                            "source": doc.name,
                            "file_path": str(doc),
                            "chunk_ids": [chunk_id(fhash, ci)],
                        }
                        dedup_refs.append((match["id"], new_ref))
                        total_chunks += 1  # count as ingested even if dedup'd
                    else:
                        document_chunks.append(dc)

                for existing_id, ref in dedup_refs:
                    update_chunk_source_refs(existing_id, ref)

                delete_file(fhash)
                stored = 0
                if document_chunks:
                    stored = store_chunks(document_chunks)
                stored += len(dedup_refs)
                _mark_ingested(str(doc), fhash)
                total_chunks += stored

                if dedup_refs:
                    progress.console.print(
                        f"  [green]  Stored {len(document_chunks)} new + {len(dedup_refs)} dedup'd chunk(s) from {doc.name}[/]"
                    )
                else:
                    progress.console.print(
                        f"  [green]  Stored {stored} chunk(s) from {doc.name}[/]"
                    )
            except Exception as e:
                progress.console.print(f"  [red]  Error processing {doc.name}: {e}[/]")
                failed += 1

            progress.advance(task)

    console.print(f"\n[bold green]Ingest complete:[/]")
    console.print(f"  Files found:     {len(docs)}")
    console.print(f"  New/changed:     {len(new_files)}")
    console.print(f"  Skipped:         {skipped}")
    console.print(f"  Failed:          {failed}")
    console.print(f"  Total chunks:    {total_chunks}")
    _sync_generation()

    return {
        "found": len(docs),
        "new": len(new_files),
        "skipped": skipped,
        "failed": failed,
        "chunks": total_chunks,
    }


# ── GitHub URL helpers ────────────────────────────────────────────────────────

_GITHUB_URL_RE = re.compile(
    r"^(https://github\.com/[\w.-]+/[\w.-]+(/.*)?$|git@github\.com:[\w.-]+/[\w.-]+(\.git)?$)"
)


def _is_github_url(value: str) -> bool:
    """Check if a string is a GitHub repository URL (HTTPS or SSH)."""
    return bool(_GITHUB_URL_RE.match(value.strip()))


def _clone_github_url(url: str, console: Console) -> Path | None:
    """Clone a GitHub repo to a temp directory. Returns the temp dir Path or None."""
    tmp = tempfile.mkdtemp(prefix="localrag_")
    try:
        console.print(f"[dim]Cloning {url} ...[/]")
        result = subprocess.run(
            ["git", "clone", "--depth", "1", url.strip(), tmp],
            capture_output=True, text=True, timeout=120,
        )
        if result.returncode != 0:
            console.print(f"[red]Git clone failed:[/] {result.stderr.strip()}")
            shutil.rmtree(tmp, ignore_errors=True)
            return None
        console.print("[green]Clone complete.[/]")
        return Path(tmp)
    except Exception as e:
        console.print(f"[red]Clone error:[/] {e}")
        shutil.rmtree(tmp, ignore_errors=True)
        return None


def ingest_repo(repo_path: str | Path) -> dict:
    """Index a code repository with code-appropriate defaults.

    Accepts a local filesystem path *or* a GitHub URL (HTTPS or SSH).  GitHub
    repos are shallow-cloned to a temp directory which is cleaned up after
    indexing.

    Uses expanded skip dirs (test dirs, build output, etc.) and code-aware
    chunking (by function/class boundaries).  The rest of the pipeline
    (embedding, storage, query) is the same as document ingestion.
    """
    from rich.console import Console

    console = Console()
    settings = get_settings()

    from local_rag.store import _get_db, _get_table
    _get_table(_get_db(), create=True)

    _check_generation(console)

    # ── Resolve input: GH URL or local path ───────────────────────────────
    _cleanup: Path | None = None
    repo_str = str(repo_path)

    if _is_github_url(repo_str):
        cloned = _clone_github_url(repo_str, console)
        if cloned is None:
            return {"found": 0, "new": 0, "skipped": 0, "failed": 0, "chunks": 0}
        resolved = cloned
        _cleanup = cloned
        display_name = repo_str
    else:
        resolved = Path(repo_str).expanduser().resolve()
        if not resolved.is_dir():
            console.print(f"[red]Not a valid directory: {resolved}[/]")
            return {"found": 0, "new": 0, "skipped": 0, "failed": 0, "chunks": 0}
        display_name = str(resolved)

    console.print(f"[bold]Indexing code repo:[/] {display_name}")

    code_files = walk_docs([resolved], extra_skip_dirs=settings.code_skip_dirs)

    if not code_files:
        console.print("[yellow]No supported code files found in repo.[/]")
        return {"found": 0, "new": 0, "skipped": 0, "failed": 0, "chunks": 0}

    # Filter to code-specific extensions only
    code_exts = {".py", ".pyi", ".rs", ".ts", ".tsx", ".js", ".go",
                 ".java", ".yaml", ".yml", ".toml", ".json", ".sql"}
    code_files = [f for f in code_files if f.suffix.lower() in code_exts]

    # Group by language for reporting
    lang_counts: dict[str, int] = {}
    for f in code_files:
        ext = f.suffix.lower()
        lang_map = {
            ".py": "Python", ".pyi": "Python",
            ".rs": "Rust",
            ".ts": "TypeScript", ".tsx": "TypeScript", ".js": "JavaScript",
            ".go": "Go",
            ".java": "Java",
            ".yaml": "YAML", ".yml": "YAML",
            ".toml": "TOML",
            ".json": "JSON",
            ".sql": "SQL",
        }
        lang = lang_map.get(ext, ext)
        lang_counts[lang] = lang_counts.get(lang, 0) + 1

    console.print(f"Found {len(code_files)} code file(s):")
    for lang, count in sorted(lang_counts.items(), key=lambda x: -x[1]):
        console.print(f"  [dim]• {lang}: {count}[/]")

    # Delegate to the shared ingest pipeline
    new_files: list[Path] = []
    skipped = 0
    failed = 0
    total_chunks = 0

    for doc in code_files:
        fhash = file_hash(doc)
        if not _is_unchanged(str(doc), fhash):
            new_files.append(doc)
        else:
            skipped += 1

    if not new_files:
        console.print("[green]All files already indexed and unchanged.[/]")
        return {"found": len(code_files), "new": 0, "skipped": skipped, "failed": 0, "chunks": 0}

    console.print(f"[bold]{len(new_files)} new/changed file(s) to process[/]")

    with Progress() as progress:
        task = progress.add_task("Indexing code...", total=len(new_files))

        for doc in new_files:
            fhash = file_hash(doc)
            try:
                progress.console.print(f"  [dim]{doc.name}[/]")
                text = extract_text(doc)
                if not text.strip():
                    failed += 1
                    progress.advance(task)
                    continue

                chunks_data = chunk_text(text, doc.name)
                if not chunks_data:
                    failed += 1
                    progress.advance(task)
                    continue

                texts_to_embed = [c[0] for c in chunks_data]
                embeddings = embed_texts(texts_to_embed)

                dedup_threshold = settings.dedup_threshold
                file_metadata = extract_file_metadata(str(doc), str(resolved))
                if file_metadata.get("type") is None:
                    file_metadata = {"type": "code", "repo": resolved.name, "file": doc.name}

                document_chunks: list[DocumentChunk] = []
                dedup_refs: list[dict] = []

                for (text_str, ci, ct), emb in zip(chunks_data, embeddings):
                    dc = DocumentChunk(
                        id=chunk_id(fhash, ci),
                        text=text_str,
                        source=doc.name,
                        file_path=str(doc),
                        chunk_index=ci,
                        total_chunks=ct,
                        file_hash=fhash,
                        embedding=emb,
                        metadata=file_metadata,
                    )
                    match = find_similar_chunk(emb, dedup_threshold, exclude_hash=fhash)
                    if match:
                        new_ref = {
                            "source": doc.name,
                            "file_path": str(doc),
                            "chunk_ids": [chunk_id(fhash, ci)],
                        }
                        dedup_refs.append((match["id"], new_ref))
                        total_chunks += 1
                    else:
                        document_chunks.append(dc)

                for existing_id, ref in dedup_refs:
                    update_chunk_source_refs(existing_id, ref)

                delete_file(fhash)
                stored = 0
                if document_chunks:
                    stored = store_chunks(document_chunks)
                stored += len(dedup_refs)
                _mark_ingested(str(doc), fhash)
                total_chunks += stored

            except Exception as e:
                progress.console.print(f"  [red]  Error: {e}[/]")
                failed += 1

            progress.advance(task)

    console.print(f"\n[bold green]Code repo ingest complete:[/]")
    console.print(f"  Files found:     {len(code_files)}")
    console.print(f"  New/changed:     {len(new_files)}")
    console.print(f"  Skipped:         {skipped}")
    console.print(f"  Failed:          {failed}")
    console.print(f"  Total chunks:    {total_chunks}")

    # Clean up temp clone if applicable
    if _cleanup is not None:
        shutil.rmtree(_cleanup, ignore_errors=True)
        console.print("[dim]Temp clone cleaned up.[/]")

    _sync_generation()

    return {
        "found": len(code_files),
        "new": len(new_files),
        "skipped": skipped,
        "failed": failed,
        "chunks": total_chunks,
    }
